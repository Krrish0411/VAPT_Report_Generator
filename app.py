# app.py - COMPLETE VAPT REPORT GENERATOR (v2)
# ============================================================
# CHANGE LOG (high level — see inline "# === MODIFIED/NEW ===" markers below
# for exact locations):
#   1. Text watermark removed; replaced with a real picture watermark
#      (faded/washed-out, centered, behind text, repeats on every page).
#   2. New "Client Name" field — flows into the header, footer context and
#      the cover page.
#   3. New "Header Logo" upload — appears top-left in the header on every page.
#   4. New "Watermark Logo" upload — used as the centered washed-out watermark.
#   5. Full 3-column page header (logo / client name / date) on every page.
#   6. Dark blue page border on every page.
#   7. Professional footer (CERT-IN line, website, Page X of Y) on every page.
#   8. Garamond enforced via style-level config (Normal + Heading 1-4 + Title)
#      instead of per-run overrides — large reduction in repeated code.
#   9. All tables: dark-blue header row / white header text / thin black
#      borders / consistent cell padding via new style_table() helper.
#  10. Risk classification, Summary-of-Findings and Vulnerability-Overview
#      tables now use full-cell severity color coding (not just colored text).
#  11. Manual Table-of-Contents replaced with a real Word TOC field
#      (TOC \o "1-3" \h \z \u) that auto-updates when the document is opened.
#  12. Cover page reworked (logo no longer duplicated in the body — it's in
#      the header already).
#  13. Vulnerability sections kept content-complete, just restructured with
#      small helpers (add_label_value, add_subsection) for less repetition
#      and consistent spacing.
#  14. Screenshots: centered, aspect-ratio preserved, spacing before/after.
#  15. Flask routes, SQLite models and existing functionality preserved.
# ============================================================

from flask import Flask, render_template, request, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
import json
import os
import re
from pathlib import Path
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import base64
from io import BytesIO

# ============================================================
# 1. CONFIGURATION
# ============================================================
app = Flask(__name__)
app.config['SECRET_KEY'] = 'vapt_secret_key_2026'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///vapt_report.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Folders
DESKTOP = Path(os.path.expanduser("~/Desktop"))
VAPT_REPORTS = DESKTOP / "VAPT_Reports"
VAPT_REPORTS.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}

# === NEW: brand / styling constants used throughout the report ===========
FONT_NAME = 'Garamond'
NAVY = "002060"                  # page border, table headers, header/footer rules
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Severity color palette (Office "standard colors" row — matches common
# enterprise VAPT report templates). Used for full-cell shading.
SEVERITY_COLORS = {
    'Critical':      'C00000',
    'High':          'FF0000',
    'Medium':        'FFC000',
    'Low':           '92D050',
    'Informational': '00B0F0',
}
# Whether each severity's background is dark enough to need white text
SEVERITY_TEXT_WHITE = {
    'Critical': True,
    'High': True,
    'Medium': False,
    'Low': False,
    'Informational': False,
}

# Default footer / cert text (overridable via the form — see /generate_report)
DEFAULT_CERT_TEXT = "CERT-In Empaneled IT Security Auditing Organization"
DEFAULT_WEBSITE = "www.protechmanize.com"
DEFAULT_ADDRESS = "www.protechmanize.com"
DEFAULT_CLASSIFICATION = "Confidential"

# === Default branding assets (header logo + watermark) =====================
# Per request, the header logo and watermark are no longer something the
# tester has to upload every time — they are ProTechmanize's own fixed
# branding, bundled with the app, and applied to EVERY generated report
# automatically. A tester can still override either one for a specific
# report via the API (header_logo / watermark_logo file fields are still
# accepted), but the UI no longer exposes an upload control for them.
# BASE_DIR makes these paths work regardless of which directory the app is
# launched from.
BASE_DIR = Path(__file__).resolve().parent
DEFAULT_HEADER_LOGO_PATH = str(BASE_DIR / "static" / "branding" / "header_logo_default.png")
DEFAULT_WATERMARK_PATH = str(BASE_DIR / "static" / "branding" / "watermark_default.png")

db = SQLAlchemy(app)

# ============================================================
# 2. DATABASE MODELS (unchanged)
# ============================================================

class VulnerabilityTemplate(db.Model):
    __tablename__ = 'vulnerability_templates'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255), nullable=False, unique=True)
    severity = db.Column(db.String(50), nullable=False)
    cwe_cve = db.Column(db.String(100))
    description = db.Column(db.Text, nullable=False)
    business_impact = db.Column(db.Text, nullable=False)
    remediation = db.Column(db.Text, nullable=False)
    references = db.Column(db.Text)

    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'severity': self.severity,
            'cwe_cve': self.cwe_cve,
            'description': self.description,
            'business_impact': self.business_impact,
            'remediation': self.remediation,
            'references': json.loads(self.references) if self.references else []
        }

# === NEW: Auto Vulnerability Knowledge Base =================================
# Every time a vulnerability is saved via /add_vulnerability, this makes sure
# it also exists in VulnerabilityTemplate (the reusable knowledge base) —
# either by creating a brand-new template (first time this vuln name is ever
# seen) or by refreshing an existing one with whatever fresh details were
# just typed in (so the knowledge base keeps improving over time instead of
# being frozen at whatever was entered the very first time).
#
# Matching is case-insensitive and trims whitespace, so "Misconfigured CSP
# Header", "misconfigured csp header", and " Misconfigured CSP Header " are
# all treated as the same template — this matters because testers won't
# always type a vulnerability name with perfectly identical casing every
# time.
def upsert_vulnerability_template(name, severity, cwe_cve, description,
                                   business_impact, remediation, references_json):
    """Creates or refreshes a VulnerabilityTemplate row from a submitted
    vulnerability. Returns (template, created: bool)."""
    clean = (name or '').strip()
    if not clean:
        return None, False

    template = VulnerabilityTemplate.query.filter(
        db.func.lower(VulnerabilityTemplate.name) == clean.lower()
    ).first()

    created = template is None
    if created:
        template = VulnerabilityTemplate(name=clean)
        db.session.add(template)

    # Only overwrite a field if the submission actually provided something —
    # this way, re-saving the same vuln with a blank "References" box (say)
    # doesn't wipe out a reference URL that was carefully filled in before.
    if severity:
        template.severity = severity
    if cwe_cve:
        template.cwe_cve = cwe_cve
    if description:
        template.description = description
    if business_impact:
        template.business_impact = business_impact
    if remediation:
        template.remediation = remediation
    if references_json:
        template.references = references_json

    return template, created

class VulnerabilityReport(db.Model):
    __tablename__ = 'vulnerability_reports'
    id = db.Column(db.Integer, primary_key=True)
    company_name = db.Column(db.String(255), nullable=False)
    name = db.Column(db.String(255), nullable=False)
    severity = db.Column(db.String(50), nullable=False)
    cwe_cve = db.Column(db.String(100))
    description = db.Column(db.Text, nullable=False)
    business_impact = db.Column(db.Text, nullable=False)
    remediation = db.Column(db.Text, nullable=False)
    references = db.Column(db.Text)
    vulnerable_urls = db.Column(db.Text)
    poc_steps = db.Column(db.Text)
    image_paths = db.Column(db.Text)
    status = db.Column(db.String(50), default='Open')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def to_dict(self):
        return {
            'id': self.id,
            'company_name': self.company_name,
            'name': self.name,
            'severity': self.severity,
            'cwe_cve': self.cwe_cve,
            'description': self.description,
            'business_impact': self.business_impact,
            'remediation': self.remediation,
            'references': json.loads(self.references) if self.references else [],
            'vulnerable_urls': json.loads(self.vulnerable_urls) if self.vulnerable_urls else [],
            'poc_steps': json.loads(self.poc_steps) if self.poc_steps else [],
            'image_paths': json.loads(self.image_paths) if self.image_paths else [],
            'status': self.status,
            'created_at': self.created_at.strftime('%d-%m-%Y')
        }

# ============================================================
# 3. VULNERABILITY TEMPLATES (unchanged)
# ============================================================

TEMPLATES = [
    {
        "name": "Misconfigured CSP Header",
        "severity": "Low",
        "cwe_cve": "CVE‑2024‑29184",
        "description": "During the assessment, it was observed that CSP is using unsafe eval and unsafe inline.",
        "business_impact": "Using unsafe-eval and unsafe-inline in CSP lets attackers run harmful scripts easily. This can lead to Cross-Site Scripting (XSS) attacks and steal user data. It weakens browser protections meant to keep the website safe.",
        "remediation": "1. Remove 'unsafe-inline' and 'unsafe-eval' from script-src.\n2. Use CSP nonce or hashes for inline scripts.\n3. Avoid wildcard sources (*) in directives.",
        "references": '["https://cheatsheetseries.owasp.org/cheatsheets/Content_Security_Policy_Cheat_Sheet.html"]'
    },
    {
        "name": "Default Credentials",
        "severity": "Critical",
        "cwe_cve": "CWE-1392",
        "description": "During our assessment, it was observed that the authentication system on the target application accepts default or weak credentials, allowing unauthorized access.",
        "business_impact": "An attacker can gain unauthorized access to sensitive areas of the application or system by using well-known default usernames and passwords. This could lead to:\n1. Data theft or modification.\n2. Privilege escalation",
        "remediation": "1. Remove or disable default accounts if not required.\n2. Force change of default passwords on first use or during setup.\n3. Implement strong password policies (complexity, length, and expiration).\n4. Enforce rate limiting and account lockout mechanisms on failed login attempts.\n5. Use multi-factor authentication (MFA) for all administrative or high-privilege accounts.",
        "references": '["https://cwe.mitre.org/data/definitions/1392.html"]'
    },
    {
        "name": "Stored Cross Site Scripting",
        "severity": "High",
        "cwe_cve": "CVE-2020-11022",
        "description": "During the assessment, it was observed that application is vulnerable to Stored Cross Site Scripting.",
        "business_impact": "Stored Cross-Site Scripting (XSS) allows attackers to inject malicious scripts into web applications, which are then permanently stored (e.g., in databases) and executed in users' browsers. This can lead to data theft, session hijacking, or full account compromise for any user who views the affected content.",
        "remediation": "1. Sanitize and validate all user input before storing it, removing or encoding dangerous characters.\n2. Use output encoding (e.g., HTML entity encoding) when displaying user data in the browser.\n3. Implement Content Security Policy (CSP) to limit script execution and reduce impact if XSS occurs.",
        "references": '["https://owasp.org/www-community/attacks/xss/"]'
    },
    {
        "name": "HTML Injection",
        "severity": "Medium",
        "cwe_cve": "CVE-2020-36475",
        "description": "During the assessment it has been found that the application is vulnerable to HTML injection attack.",
        "business_impact": "Attackers can inject malicious HTML code into the application, altering the content displayed to users. facilitating cross-site scripting attacks and compromising web security.",
        "remediation": "1. Validate and sanitize all user inputs to prevent HTML code from being injected.\n2. Encode outputs to ensure that any HTML tags are rendered harmless.\n3. Implement CSP to limit the sources from which content can be loaded and executed.",
        "references": '["https://cwe.mitre.org/data/definitions/80.html"]'
    },
    {
        "name": "Clickjacking",
        "severity": "Informational",
        "cwe_cve": "CWE-1021",
        "description": "During the assessment it was observed that the application does not implement clickjacking protection mechanisms such as X-Frame-Options or Content-Security-Policy: frame-ancestors, allowing the page to be embedded within an iframe.",
        "business_impact": "The absence of frame-busting protections could allow malicious websites to trick users into unintended actions.",
        "remediation": "Implement a restrictive Content-Security-Policy with frame-ancestors 'none' to prevent unauthorized framing.",
        "references": '["https://cwe.mitre.org/data/definitions/1021.html"]'
    },
    {
        "name": "Cross Origin Resource Sharing",
        "severity": "Informational",
        "cwe_cve": "CWE-942",
        "description": "During the assessment it was observed that the CORS policy is configured to reflect arbitrary Origin values and allows cross-origin requests with credentials enabled.",
        "business_impact": "An overly permissive CORS policy combined with credentialed requests could allow unauthorized cross-origin access if future endpoints return sensitive or user-specific information.",
        "remediation": "1. Restrict Access-Control-Allow-Origin to trusted domains rather than reflecting arbitrary values.\n2. Avoid using Access-Control-Allow-Credentials: true on endpoints that do not require credentials.",
        "references": '["https://cwe.mitre.org/data/definitions/942.html"]'
    },
    {
        "name": "Missing Security Headers",
        "severity": "Informational",
        "cwe_cve": "CWE-1021",
        "description": "The server does not implement several recommended HTTP security headers, such as Strict-Transport-Security (HSTS), X-Content-Type-Options, Content-Security-Policy.",
        "business_impact": "Missing security headers can increase exposure to web attacks such as clickjacking, MIME-type sniffing, or information disclosure if the server functionality changes in the future.",
        "remediation": "Implement recommended security headers:\n1. Strict-Transport-Security: max-age=31536000; includeSubDomains\n2. X-Content-Type-Options: nosniff\n3. Content-Security-Policy:\n   - frame-ancestors 'none'\n   - default-src 'self'; script-src 'self' https://trusted-cdn.com 'nonce-abc123';\n   - style-src 'self'; img-src 'self' data:;\n   - object-src 'none'; base-uri 'self';\n   - require-trusted-types-for 'script';",
        "references": '["https://cheatsheetseries.owasp.org/cheatsheets/HTTP_Headers_Cheat_Sheet.html"]'
    },
    {
        "name": "Weak Ciphers Enabled",
        "severity": "Low",
        "cwe_cve": "CWE-326",
        "description": "During the assessment, it was observed that application supports weak cryptographic ciphers like CBC (Cipher Block Chaining).",
        "business_impact": "Using weak CBC ciphers increases the risk of data interception and decryption, potentially exposing sensitive business information during transmission and weakening overall encryption strength.",
        "remediation": "Disable weak CBC ciphers and enable stronger suites like AES-GCM to ensure secure and modern encryption aligned with industry best practices.",
        "references": '["https://owasp.org/www-project-web-security-testing-guide/v41/4-Web_Application_Security_Testing/09-Testing_for_Weak_Cryptography/01-Testing_for_Weak_SSL_TLS_Ciphers_Insufficient_Transport_Layer_Protection"]'
    },
    {
        "name": "Improper Error Handling",
        "severity": "Low",
        "cwe_cve": "CWE-209",
        "description": "The application fails to properly handle invalid requests. Instead of returning a generic error page, the server displays a detailed ASP.NET runtime error page.",
        "business_impact": "Detailed error responses reveal information about the application's framework and configuration settings. This information may assist an attacker in understanding the application's architecture and facilitate the identification and exploitation of additional vulnerabilities.",
        "remediation": "Implement custom error pages, disable verbose error messages, and provide only generic error responses.",
        "references": '["https://cheatsheetseries.owasp.org/cheatsheets/Error_Handling_Cheat_Sheet.html"]'
    },
    {
        "name": "Unrestricted Internal Service Access",
        "severity": "Informational",
        "cwe_cve": "N/A",
        "description": "During the assessment, TCP port 1433 was identified as open.",
        "business_impact": "An attacker may be able to identify and interact with the exposed database service. This exposure increases the likelihood of unauthorized access attempts and could facilitate compromise of the database environment if additional weaknesses are present.",
        "remediation": "Restrict access to TCP port 1433 using firewall rules or access controls so it is only reachable from trusted hosts. If not required, close the port.",
        "references": '[]'
    }
]

# ============================================================
# 4. FOLDER MANAGER
# ============================================================

def clean_name(name):
    clean = name.replace(' ', '_')
    clean = re.sub(r'[^a-zA-Z0-9_\-]', '', clean)
    return clean

def get_vulnerability_folder(company_name, vuln_name):
    company_path = VAPT_REPORTS / clean_name(company_name)
    company_path.mkdir(parents=True, exist_ok=True)
    vuln_path = company_path / clean_name(vuln_name)
    vuln_path.mkdir(parents=True, exist_ok=True)
    return vuln_path

def save_poc_image(company_name, vuln_name, image_file):
    vuln_path = get_vulnerability_folder(company_name, vuln_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ext = os.path.splitext(image_file.filename)[1]
    image_name = f"poc_{timestamp}{ext}"
    image_path = vuln_path / image_name
    image_file.save(str(image_path))
    return str(image_path)

def get_images(company_name, vuln_name):
    vuln_path = get_vulnerability_folder(company_name, vuln_name)
    images = []
    if vuln_path.exists():
        for file in vuln_path.iterdir():
            if file.suffix.lower() in {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}:
                images.append(str(file))
    return sorted(images)

# === NEW: persist / reuse the header & watermark logo per company =========
def save_or_reuse_asset(company_name, file_obj, kind, default_path=None):
    """Saves a newly uploaded header/watermark logo under the company's
    folder so the tester only has to upload it once per engagement (only
    relevant if they ever choose to override the default branding). If no
    new file is uploaded on this request, falls back to:
      1. whatever was saved previously for this company (if anything), then
      2. `default_path` — ProTechmanize's bundled default logo/watermark,
         used automatically so nothing has to be uploaded for the normal
         case.
    Returns an absolute path string, or None if nothing is available at all.

    `kind` is a short slug, e.g. 'header_logo' or 'watermark_logo'.
    """
    assets_dir = VAPT_REPORTS / clean_name(company_name) / "_assets"
    assets_dir.mkdir(parents=True, exist_ok=True)

    if file_obj and getattr(file_obj, 'filename', ''):
        ext = os.path.splitext(file_obj.filename)[1].lower().lstrip('.')
        if ext not in ALLOWED_EXTENSIONS:
            ext = 'png'
        dest = assets_dir / f"{kind}.{ext}"
        # Remove any previous asset with a different extension so we don't
        # accidentally keep using a stale logo.
        for old_ext in ALLOWED_EXTENSIONS:
            old_file = assets_dir / f"{kind}.{old_ext}"
            if old_file.exists() and old_file != dest:
                old_file.unlink()
        file_obj.save(str(dest))
        return str(dest)

    for ext in ALLOWED_EXTENSIONS:
        candidate = assets_dir / f"{kind}.{ext}"
        if candidate.exists():
            return str(candidate)

    if default_path and os.path.exists(default_path):
        return default_path

    return None

# ============================================================
# 5. STYLING HELPERS  (=== NEW SECTION ===)
# These replace the old, ad-hoc per-call font loops and the text watermark.
# Everything here was validated against a rendered prototype before being
# wired into generate_vapt_report().
# ============================================================

def style_run(run, size=11, bold=False, italic=False, color=None, font_name=FONT_NAME):
    """One-liner replacement for the repeated
    `run.font.name = ...; run.font.size = ...; run.bold = ...` blocks."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run

def configure_heading_style(doc, style_name, size_pt, bold=True, remove_border=False,
                             add_bottom_border=False):
    """Configures a built-in style (Title / Heading 1-4) once, so every
    doc.add_heading() call automatically comes out in Garamond/bold/black
    without needing per-call run loops. Also strips the blue color and the
    bottom border Word's default 'Title' style ships with.

    add_bottom_border=True adds a thin gray rule below the heading text —
    matches the reference report's "4.2.x" finding-title headings, which
    have a thin gray line running under them."""
    st = doc.styles[style_name]
    st.font.name = FONT_NAME
    st.font.size = Pt(size_pt)
    st.font.bold = bold
    st.font.italic = False   # Word's built-in 'Heading 4' is italic by default; override it
    st.font.color.rgb = BLACK
    rPr = st.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if remove_border or add_bottom_border:
        pPr = st.element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            st.element.insert(0, pPr)
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)
        if add_bottom_border:
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), '808080')
            pBdr.append(bottom)
            pPr.append(pBdr)
            space = OxmlElement('w:spacing')
            space.set(qn('w:after'), '120')
            pPr.append(space)

def set_cell_border(cell, sz=4, color="000000"):
    """Thin single border on all four sides of a cell (0.5pt by default)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def clear_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    """Consistent cell padding (in twentieths of a point / dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)

def shade_cell(cell, hex_color):
    """Full cell background fill."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_cell_text(cell, size=10, bold=False, color=None, align=None, font_name=FONT_NAME):
    """Applies font styling to whatever runs already exist in a cell. Must be
    called AFTER the cell's text is set (this is the bug we fixed from the
    original code, which styled cells before any runs existed)."""
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        if not paragraph.runs:
            continue
        for run in paragraph.runs:
            style_run(run, size=size, bold=bold, color=color, font_name=font_name)

def prevent_row_split(table):
    """Stops a table row from being torn across a page boundary (the
    'Critical/High/.../Informational' rows in the Risk Classification table
    can otherwise split mid-row right at a page break)."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))

def style_table(table, header_rows=1, header_bg=NAVY, body_font_size=10,
                 header_font_size=10, border_color="000000"):
    """Applies the standard report table look to every cell: thin black
    borders, consistent padding, and (for the header row(s)) a dark-blue
    fill with white bold text. Call this AFTER all cell text has been set."""
    for r_idx, row in enumerate(table.rows):
        is_header = r_idx < header_rows
        for cell in row.cells:
            set_cell_border(cell, color=border_color)
            set_cell_margins(cell)
            if is_header:
                shade_cell(cell, header_bg)
                style_cell_text(cell, size=header_font_size, bold=True, color=WHITE)
            else:
                style_cell_text(cell, size=body_font_size, bold=False, color=BLACK)
    prevent_row_split(table)

def add_page_border(doc, color=NAVY, size=18, space=24):
    """Single dark-blue border around every page (offset from the page
    edge), matching enterprise report templates."""
    for section in doc.sections:
        sectPr = section._sectPr
        for old in sectPr.findall(qn('w:pgBorders')):
            sectPr.remove(old)
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for edge in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(size))
            el.set(qn('w:space'), str(space))
            el.set(qn('w:color'), color)
            pgBorders.append(el)
        sectPr.append(pgBorders)

def add_picture_watermark(section, image_path):
    """=== NEW: replaces the old text-based add_watermark() ===
    Inserts a faded/washed-out picture watermark into the header so it
    repeats on every page, using the same VML mechanism Word's own
    Insert > Watermark > Picture feature generates (gain/blacklevel
    'washout' on v:imagedata — this is the technique that actually renders
    a faded image reliably across Word and LibreOffice; the modern
    DrawingML lumMod/lumOff recolor was tried first but several renderers,
    including LibreOffice, ignore it and show the image at full strength).

    Centering uses VML's own mso-position-horizontal:center /
    mso-position-vertical:center keywords relative to the page, instead of
    computing a manual left/top pixel offset — that manual math is what
    caused the watermark to drift off-center for some page/image
    combinations. The 'center' keyword always centers correctly regardless
    of the image's size or aspect ratio."""
    if not image_path or not os.path.exists(image_path):
        return
    header = section.header
    header_part = header.part
    rId, image = header_part.get_or_add_image(image_path)

    # Size relative to the page; height follows the image's own aspect
    # ratio so it's never stretched or squashed.
    page_w_pt = section.page_width.pt
    target_w_pt = page_w_pt * 0.45
    aspect = image.height / float(image.width)
    target_h_pt = target_w_pt * aspect

    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = paragraph.add_run()

    pict_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<v:shapetype id="_x0000_t75" coordsize="21600,21600" o:spt="75" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" filled="f" stroked="f" '
        'o:preferrelative="t" path="m@4@5l@4@11@9@11@9@5xe">'
        '<v:stroke joinstyle="miter"/>'
        '<v:formulas>'
        '<v:f eqn="if lineDrawn pixelLineWidth 0"/><v:f eqn="sum @0 1 0"/><v:f eqn="sum 0 0 @1"/>'
        '<v:f eqn="prod @2 1 2"/><v:f eqn="prod @3 21600 pixelWidth"/><v:f eqn="prod @3 21600 pixelHeight"/>'
        '<v:f eqn="sum @0 0 1"/><v:f eqn="prod @6 1 2"/><v:f eqn="prod @7 21600 pixelWidth"/>'
        '<v:f eqn="sum @8 21600 0"/><v:f eqn="prod @7 21600 pixelHeight"/><v:f eqn="sum @10 21600 0"/>'
        '</v:formulas>'
        '<v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="rect"/>'
        '<o:lock v:ext="edit" aspectratio="t"/>'
        '</v:shapetype>'
        f'<v:shape id="WatermarkPicture01" o:spid="_x0000_s2050" type="#_x0000_t75" '
        f'style="position:absolute;margin-left:0;margin-top:0;width:{target_w_pt:.2f}pt;'
        f'height:{target_h_pt:.2f}pt;z-index:-251650048;'
        f'mso-position-horizontal:center;mso-position-horizontal-relative:page;'
        f'mso-position-vertical:center;mso-position-vertical-relative:page" '
        'o:allowincell="f" xmlns:o="urn:schemas-microsoft-com:office:office">'
        f'<v:imagedata r:id="{rId}" o:title="watermark" gain="19661f" blacklevel="22938f"/>'
        '<w10:wrap xmlns:w10="urn:schemas-microsoft-com:office:word" anchorx="page" anchory="page"/>'
        '</v:shape>'
        '</w:pict>'
    )
    run._r.append(etree.fromstring(pict_xml.encode('utf-8')))

def build_page_header(section, logo_path, client_name, report_date):
    """3-column single-line letterhead header, matching the reference
    exactly:

        [Logo]        Client Name (centered)        Report Date (right)

    Built as a borderless 1-row/3-column table so all three pieces sit on
    ONE line regardless of page width — a plain paragraph with a hanging
    indent (the previous approach) let the date wrap to a second line for
    longer client names/logos, which is exactly the bug being fixed here."""
    header = section.header
    header.is_linked_to_previous = False
    header_p = header.paragraphs[0]
    header_p.text = ""

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = header.add_table(rows=1, cols=3, width=usable_width)
    table.autofit = False
    col_widths = [Emu(int(usable_width * 0.22)), Emu(int(usable_width * 0.56)), Emu(int(usable_width * 0.22))]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
        table.rows[0].cells[i].width = width

    logo_cell, name_cell, date_cell = table.rows[0].cells
    for cell in (logo_cell, name_cell, date_cell):
        clear_cell_borders(cell)
        set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Left: logo
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path and os.path.exists(logo_path):
        run = lp.add_run()
        try:
            run.add_picture(logo_path, height=Pt(30))
        except Exception:
            pass

    # Center: client name
    np = name_cell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(np.add_run(client_name or ''), size=13, bold=False)

    # Right: report date
    dp = date_cell.paragraphs[0]
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(dp.add_run(report_date or ''), size=13, bold=True)

    # Thin navy rule separating the header from the body content, directly
    # under the header table
    rule_p = header.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(4)
    pPr2 = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), NAVY)
    pBdr.append(bottom)
    pPr2.append(pBdr)

def _add_page_number_field(paragraph, font_size=9, color=None, bold=False):
    """Inserts 'Page <PAGE> of <NUMPAGES>' as live Word fields."""
    def field(instr):
        r = paragraph.add_run()
        style_run(r, size=font_size, color=color, bold=bold)
        begin = OxmlElement('w:fldChar')
        begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instr
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        r._r.append(begin)
        r._r.append(instrText)
        r._r.append(end)

    style_run(paragraph.add_run("Page "), size=font_size, color=color, bold=bold)
    field("PAGE")
    style_run(paragraph.add_run(" of "), size=font_size, color=color, bold=bold)
    field("NUMPAGES")

LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)

def add_bottom_navy_bar(section, height_pt=28):
    """Decorative solid navy (#002060) bar sitting at the very bottom edge
    of the page, BELOW the footer text — this is a separate design element
    from the footer content, matching the reference report (a plain
    #002060 filled rectangle anchored to the bottom of the page, full
    page width, no text in it, no border)."""
    footer = section.footer
    p = footer.paragraphs[0] if not footer.paragraphs[0].runs and not footer.paragraphs[0].text else footer.add_paragraph()
    run = p.add_run()
    page_w_pt = section.page_width.pt
    rect_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office">'
        f'<v:rect id="BottomNavyBar" o:spid="_x0000_s9999" '
        f'style="position:absolute;margin-left:0;margin-top:0;width:{page_w_pt:.2f}pt;'
        f'height:{height_pt}pt;z-index:-251650000;'
        f'mso-position-horizontal:center;mso-position-horizontal-relative:page;'
        f'mso-position-vertical:bottom;mso-position-vertical-relative:page" '
        'fillcolor="#002060" stroked="f" o:allowincell="f">'
        '<w10:wrap xmlns:w10="urn:schemas-microsoft-com:office:word" anchorx="page" anchory="page"/>'
        '</v:rect>'
        '</w:pict>'
    )
    run._r.append(etree.fromstring(rect_xml.encode('utf-8')))

def build_page_footer(section, cert_text, website_text, classification_text="Confidential",
                       address_text=None):
    """=== Matches the reference report's footer exactly ===
    Plain white background, black text (NOT a shaded navy band — that navy
    color in the reference is a separate decorative bar at the very bottom
    edge of the page, added by add_bottom_navy_bar(), not a fill behind the
    footer text):
      1. "CERT-IN Empaneled" (bold, black), with a thin rule underneath.
      2. One line: "Classification: <classification_text>   <address>"
         on the left, "Page X of Y" on the right — plain black text.
      3. A centered hyperlink-styled line with the company website
         (blue, underlined), matching the reference.
    Renders on every page (single section, no odd/even or first-page
    footer overrides)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    usable_width = section.page_width - section.left_margin - section.right_margin

    # --- Line 1: "CERT-IN Empaneled" + thin rule ---
    p = footer.paragraphs[0]
    p.text = ""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_rule = OxmlElement('w:bottom')
    bottom_rule.set(qn('w:val'), 'single')
    bottom_rule.set(qn('w:sz'), '8')
    bottom_rule.set(qn('w:space'), '4')
    bottom_rule.set(qn('w:color'), '8EA9DB')
    pBdr.append(bottom_rule)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(cert_text), size=9, bold=True, color=BLACK)

    # --- Line 2: Classification + address (left)  /  Page X of Y (right) ---
    table = footer.add_table(rows=1, cols=2, width=usable_width)
    table.autofit = False
    col_widths = [Emu(int(usable_width * 0.8)), Emu(int(usable_width * 0.2))]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
        table.rows[0].cells[i].width = width
    left_cell, right_cell = table.rows[0].cells
    for cell in (left_cell, right_cell):
        clear_cell_borders(cell)
        set_cell_margins(cell, top=20, bottom=0, left=0, right=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(lp.add_run(f'Classification: {classification_text}'), size=8, bold=False, color=BLACK)
    style_run(lp.add_run('      ' + (address_text or '')), size=8, bold=False, color=BLACK)

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_number_field(rp, font_size=9, color=BLACK, bold=True)

    # --- Line 3: centered, hyperlink-styled website ---
    if website_text:
        wp = footer.add_paragraph()
        wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wp.paragraph_format.space_before = Pt(0)
        link_run = wp.add_run(website_text)
        style_run(link_run, size=8, color=LINK_BLUE)
        link_run.underline = True

    # --- Decorative navy bar at the very bottom edge of the page ---
    add_bottom_navy_bar(section)

def insert_toc_field(doc):
    """=== NEW: replaces the manual, hand-built Table of Contents ===
    Inserts a real Word TOC field (TOC \\o "1-3" \\h \\z \\u) covering
    Heading levels 1-3, with hyperlinks and a hidden tab leader — exactly
    what Word's Reference > Table of Contents button inserts. Combined with
    enable_auto_update_fields(), Word will populate/refresh it the moment
    the document is opened, no manual right-click required."""
    paragraph = doc.add_paragraph()

    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(fld_begin)
    run._r.append(instrText)

    sep_run = paragraph.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    sep_run._r.append(fld_sep)

    placeholder = paragraph.add_run(
        "Right-click here and choose \u201cUpdate Field\u201d (or press F9) "
        "to generate the Table of Contents."
    )
    style_run(placeholder, italic=True, size=10)

    end_run = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    end_run._r.append(fld_end)
    return paragraph

def enable_auto_update_fields(doc):
    """Forces Word to refresh all fields (including the TOC) automatically
    when the document is opened, per the brief's 'TOC should update when
    Word opens' requirement."""
    settings = doc.settings.element
    for old in settings.findall(qn('w:updateFields')):
        settings.remove(old)
    el = OxmlElement('w:updateFields')
    el.set(qn('w:val'), 'true')
    settings.append(el)

# ============================================================
# === NEW: small content helpers used inside generate_vapt_report() ===
# These remove the repeated 4-6 line blocks the original code used for
# every label/value line and every Description:/Remediation:/etc. block.
# ============================================================

def add_label_value(doc, label, value, value_bold=True, value_color=None, label_bold=True):
    """e.g. 'Severity Level:  <tab><tab>  Critical' with a bold label and a
    styled value — reference uses two tabs between label and value."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    style_run(p.add_run(f'{label}: '), bold=label_bold)
    style_run(p.add_run('\t\t'), bold=label_bold)
    style_run(p.add_run(value or ''), bold=value_bold, color=value_color)
    return p

def add_subsection(doc, title, body_text, heading_level=4):
    """A 'Description:' / 'Business Impact:' / 'Remediation:' style block.
    Matches the reference report: the label is bold and sits inline with
    the body text in the SAME paragraph (not a separate heading line),
    and the body is justified rather than left-aligned."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_run(p.add_run(f'{title}: '), bold=True)
    style_run(p.add_run(body_text or ''))
    return p

def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.line_spacing = 1.15
        style_run(p.add_run(item))

def add_screenshot(doc, image_path, width_in=6.0):
    """=== Screenshot formatting per the brief ===
    Centered, aspect ratio preserved (python-docx scales height
    automatically when only width is given), with breathing room before
    and after."""
    if not image_path or not os.path.exists(image_path):
        return
    p_before = doc.add_paragraph()
    p_before.paragraph_format.space_after = Pt(4)
    try:
        doc.add_picture(image_path, width=Inches(width_in))
    except Exception:
        return
    pic_paragraph = doc.paragraphs[-1]
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_paragraph.paragraph_format.space_after = Pt(10)

def severity_color(severity):
    return SEVERITY_COLORS.get(severity)

def severity_text_color(severity):
    return WHITE if SEVERITY_TEXT_WHITE.get(severity, False) else BLACK

# ============================================================
# 6. REPORT GENERATOR  (=== HEAVILY MODIFIED ===)
# Same overall section order as the original (cover -> TOC -> document
# control -> executive summary -> summary of findings -> vulnerability
# analysis -> OWASP top 10 -> appendix), now built with the helpers above.
# ============================================================

def generate_vapt_report(company_name, scope_urls, tester_name, reviewer_name, approver_name,
                          vulnerabilities, client_name=None, header_logo_path=None,
                          watermark_logo_path=None, report_date=None,
                          cert_in_text=DEFAULT_CERT_TEXT, company_website=DEFAULT_WEBSITE,
                          classification_text="Confidential", address_text=None):
    """Generate the VAPT report.

    New optional parameters (all backward compatible — existing callers that
    only pass the original positional arguments keep working unchanged):
      client_name        -> shown in the page header & cover page
                             (falls back to company_name if omitted)
      header_logo_path   -> top-left logo on every page (defaults to the
                             bundled ProTechmanize logo if not given)
      watermark_logo_path-> centered faded watermark on every page (defaults
                             to the bundled ProTechmanize watermark)
      report_date         -> shown in the header, under the client name
                             (defaults to today)
      cert_in_text/company_website -> footer text
    """
    doc = Document()
    client_name = client_name or company_name
    report_date = report_date or datetime.now().strftime('%d-%m-%Y')
    header_logo_path = header_logo_path or DEFAULT_HEADER_LOGO_PATH
    watermark_logo_path = watermark_logo_path or DEFAULT_WATERMARK_PATH

    # ===== Page margins (matches the enterprise template's proportions) =====
    section = doc.sections[0]
    section.top_margin = Pt(72)       # 1"
    section.bottom_margin = Pt(40.5)  # 0.5625"
    section.left_margin = Pt(72)      # 1"
    section.right_margin = Pt(54)     # 0.75"
    section.header_distance = Pt(43.2)
    section.footer_distance = Pt(28.8)

    # ===== Garamond everywhere, configured once at the style level so every
    # doc.add_heading()/doc.add_paragraph() call inherits it automatically
    # (this replaces the dozens of repeated per-run font loops from the
    # original code) =====
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    configure_heading_style(doc, 'Title', 14, remove_border=True)
    configure_heading_style(doc, 'Heading 1', 16)
    configure_heading_style(doc, 'Heading 2', 14)
    configure_heading_style(doc, 'Heading 3', 12, add_bottom_border=True)
    configure_heading_style(doc, 'Heading 4', 11)

    # ===== Page border / header / footer / watermark / TOC auto-update =====
    add_page_border(doc)
    build_page_header(section, header_logo_path, client_name, report_date)
    build_page_footer(section, cert_in_text, company_website,
                       classification_text=classification_text, address_text=address_text)
    add_picture_watermark(section, watermark_logo_path)
    enable_auto_update_fields(doc)

    # ============================================================
    # COVER PAGE  (=== MODIFIED: logo removed from body — it's already
    # in the header on every page, including this one ===)
    # ============================================================
    # Generous space above the title, like the reference cover page
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer_run = spacer.add_run()
    spacer_run.font.size = Pt(1)
    for _ in range(8):
        spacer.add_run().add_break()

    title = doc.add_heading('Web Application Penetration Testing Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run(f'Organization: {client_name}'), size=16, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run('Document Timestamp'), bold=True)

    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Version', 'Issued by', 'Reviewed by', 'Issued Date']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    data = ['1.0', tester_name, reviewer_name, report_date]
    for i, val in enumerate(data):
        table.rows[1].cells[i].text = val
        table.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table)

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS  (=== MODIFIED: real Word TOC field, not a manual
    # hand-built list of dots and page numbers ===)
    # ============================================================
    doc.add_heading('Table of Contents', level=1)
    insert_toc_field(doc)
    doc.add_page_break()

    # ============================================================
    # 1. DOCUMENT CONTROL
    # ============================================================
    doc.add_heading('1. Document Control', level=1)

    doc.add_heading('1.1 Document Approvers', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run(
        "The individuals who reviewed, verified and approved the content of this "
        "document based on the requirements of the defined services during this "
        "document's creation are listed in the table below:"
    ))

    table = doc.add_table(rows=2, cols=3)
    for i, header in enumerate(['Approver Name', 'Approver Designation', 'Date Content was Approved']):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table.rows[1]
    row.cells[0].text = approver_name
    row.cells[1].text = 'Manager - Cyber Security'
    row.cells[2].text = report_date
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table)

    doc.add_heading('1.2 Amendment Record', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run('The following Change History log contains a record of changes made to this document.'))

    table = doc.add_table(rows=2, cols=4)
    for i, header in enumerate(['Version', 'Issuer', 'Section / Nature of Change', 'Published/Revised Date']):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table.rows[1]
    row.cells[0].text = '1.0'
    row.cells[1].text = tester_name
    row.cells[2].text = 'Initial VAPT'
    row.cells[3].text = report_date
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table)

    doc.add_page_break()

    # ============================================================
    # 2. EXECUTIVE SUMMARY
    # ============================================================
    doc.add_heading('2. Executive Summary', level=1)

    doc.add_heading('2.1 Scope of Testing', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run('The following URLs were in scope and included in this assessment:'))

    table = doc.add_table(rows=1 + len(scope_urls), cols=5)
    for i, header in enumerate(['Sr. No.', 'Application Name', 'Tested URL', 'Tested Environment', 'Testing Methodology']):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, url_info in enumerate(scope_urls, 1):
        row = table.rows[idx]
        row.cells[0].text = str(idx)
        row.cells[1].text = url_info.get('name', 'Application')
        row.cells[2].text = url_info.get('url', '')
        row.cells[3].text = url_info.get('env', 'UAT/PROD')
        row.cells[4].text = url_info.get('methodology', 'BlackBox')
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table, body_font_size=9, header_font_size=9)

    doc.add_heading('2.2 Milestone', level=2)
    table = doc.add_table(rows=3, cols=2)
    for i, row_data in enumerate([('Milestone / Deliverable', 'Date'),
                                   ('VAPT Initiation', report_date),
                                   ('VAPT Completion', report_date)]):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        for cell in table.rows[i].cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table)

    doc.add_heading('2.3 Risk Classification', level=2)
    # === MODIFIED: real 2-column table with full-cell severity color
    # coding (label cell colored, description cell plain white) instead of
    # the original's plain colored-text paragraphs ===
    risk_rows = [
        ('Critical', 'These vulnerabilities can allow attackers to take complete control of '
                     'your web applications and web servers. In exploiting this type of '
                     'vulnerability, attackers could carry out a range of malicious acts.'),
        ('High', 'A High severity vulnerability, which means that on exploiting such '
                 'vulnerabilities, attackers can view information about your system that '
                 'helps them find or exploit other vulnerabilities that enable them to '
                 'access sensitive user and administrator information.'),
        ('Medium', 'Potential weakness in controls, which could develop into an exposure. '
                   'Or issues that represent areas of concern and may impact controls. '
                   'They should be addressed reasonably promptly.'),
        ('Low', 'Potential weaknesses in controls, which in combination with other '
                'weaknesses can develop into exposure. Suggested improvements not '
                'immediately/directly affecting controls.'),
        ('Informational', 'Weaknesses mentioned under these sections are informational and '
                           'are best practices. Either these weaknesses cannot be exploited '
                           'directly or are very difficult to exploit due to multiple constraints.'),
    ]
    total_w = section.page_width - section.left_margin - section.right_margin
    label_w, desc_w = Emu(int(total_w * 0.18)), Emu(int(total_w * 0.82))
    risk_table = doc.add_table(rows=len(risk_rows), cols=2)
    risk_table.autofit = False
    risk_table.columns[0].width = label_w
    risk_table.columns[1].width = desc_w
    for i, (level, desc) in enumerate(risk_rows):
        label_cell, desc_cell = risk_table.rows[i].cells
        label_cell.width, desc_cell.width = label_w, desc_w
        label_cell.text = f'{level} Risk'
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(label_cell, severity_color(level))
        style_cell_text(label_cell, size=11, bold=True, color=severity_text_color(level),
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        desc_cell.text = desc
        style_cell_text(desc_cell, size=10, color=BLACK)
        for cell in (label_cell, desc_cell):
            set_cell_border(cell)
            set_cell_margins(cell)
    prevent_row_split(risk_table)

    doc.add_page_break()

    # ============================================================
    # 3. SUMMARY OF FINDINGS
    # ============================================================
    doc.add_heading('3. Summary of Findings', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run('This section lists and bifurcates the total number of vulnerabilities reported, based on risk level.'))

    stats = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0, 'Informational': 0}
    for v in vulnerabilities:
        if v.get('severity', '') in stats:
            stats[v.get('severity', '')] += 1

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run('The vulnerability assessment and penetration testing were performed on '))
    style_run(p.add_run(f'{client_name} '), bold=True)
    style_run(p.add_run('applications. The test revealed '))
    summary_parts = [f'{stats[s]} {s}' for s in ['Critical', 'High', 'Medium', 'Low', 'Informational'] if stats[s] > 0]
    style_run(p.add_run(', '.join(summary_parts)), bold=True)
    style_run(p.add_run(
        ' vulnerabilities on the scope items. We recommend implementing the measures '
        'suggested for mitigation for each finding to improve the security posture of '
        'the affected systems.'
    ))

    # === MODIFIED: header row uses full severity-color cell shading
    # (matches the sample) ===
    table = doc.add_table(rows=2, cols=6)
    sev_order = ['Critical', 'High', 'Medium', 'Low', 'Informational']
    for i, header in enumerate(sev_order + ['Total']):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header in SEVERITY_COLORS:
            shade_cell(cell, severity_color(header))
            style_cell_text(cell, size=11, bold=True, color=severity_text_color(header),
                             align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            shade_cell(cell, "BFBFBF")
            style_cell_text(cell, size=11, bold=True, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_border(cell)
        set_cell_margins(cell)
    row = table.rows[1]
    for i, sev in enumerate(sev_order):
        row.cells[i].text = str(stats[sev])
    row.cells[5].text = str(sum(stats.values()))
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)
        set_cell_margins(cell)
        style_cell_text(cell, size=11, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ============================================================
    # 4. VULNERABILITY ANALYSIS & FINDINGS
    # ============================================================
    doc.add_heading('4. Vulnerability Analysis & Findings', level=1)

    doc.add_heading('4.1 Vulnerability Overview', level=2)
    table = doc.add_table(rows=1 + len(vulnerabilities), cols=5)
    for i, header in enumerate(['Sr. No', 'Vulnerability', 'Severity', 'CWE/CVE', 'Status']):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, vuln in enumerate(vulnerabilities, 1):
        row = table.rows[idx]
        row.cells[0].text = str(idx)
        row.cells[1].text = vuln.get('name', '')
        row.cells[2].text = vuln.get('severity', '')
        row.cells[3].text = vuln.get('cwe_cve', '')
        row.cells[4].text = vuln.get('status', 'Open')
        for c in (0, 2, 3, 4):
            row.cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table)
    # Color the Severity column per-row, full-cell, after the base styling
    # pass above so it overrides the plain white fill on that column only.
    for idx, vuln in enumerate(vulnerabilities, 1):
        sev = vuln.get('severity', '')
        if sev in SEVERITY_COLORS:
            cell = table.rows[idx].cells[2]
            shade_cell(cell, severity_color(sev))
            style_cell_text(cell, size=10, bold=True, color=severity_text_color(sev),
                             align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    doc.add_heading('4.2 Detailed Findings', level=2)

    for idx, vuln in enumerate(vulnerabilities, 1):
        # === MODIFIED: numbered "4.2.<n> <name>" to match the sample's
        # hierarchical numbering (and so the TOC field nests it correctly
        # under "4.2 Detailed Findings") ===
        doc.add_heading(f'4.2.{idx} {vuln.get("name", "")}', level=3)

        sev = vuln.get('severity', '')
        sev_rgb = RGBColor.from_string(severity_color(sev)) if sev in SEVERITY_COLORS else None
        add_label_value(doc, 'Severity Level', sev, value_color=sev_rgb)
        status_val = vuln.get('status', 'Open')
        status_color = RGBColor(0xEE, 0x00, 0x00) if status_val.strip().lower() == 'open' else None
        add_label_value(doc, 'Status', status_val, value_color=status_color)
        add_label_value(doc, 'CVE/CWE', vuln.get('cwe_cve', ''))

        add_subsection(doc, 'Description', vuln.get('description', ''))
        add_subsection(doc, 'Business Impact', vuln.get('business_impact', ''))

        vuln_urls = vuln.get('vulnerable_urls', [])
        if isinstance(vuln_urls, str):
            vuln_urls = [vuln_urls] if vuln_urls.strip() else []
        if vuln_urls:
            up = doc.add_paragraph()
            up.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            up.paragraph_format.space_before = Pt(8)
            up.paragraph_format.space_after = Pt(8)
            style_run(up.add_run('Vulnerable URLs: '), bold=True)
            for i, url in enumerate(vuln_urls):
                if i > 0:
                    up.add_run().add_break()
                style_run(up.add_run(url))

        poc_steps = vuln.get('poc_steps', [])
        if poc_steps:
            for step in poc_steps:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                style_run(p.add_run(f'POC{step.get("number", "")}: '), bold=True)
                style_run(p.add_run(step.get('text', '')))
                add_screenshot(doc, step.get('image', ''))
        else:
            p = doc.add_paragraph()
            style_run(p.add_run('No POC steps provided.'), italic=True)

        add_subsection(doc, 'Remediation', vuln.get('remediation', ''))

        refs = vuln.get('references', [])
        if isinstance(refs, str):
            refs = [refs] if refs.strip() else []
        if refs:
            rp = doc.add_paragraph()
            rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            rp.paragraph_format.space_before = Pt(8)
            rp.paragraph_format.space_after = Pt(8)
            style_run(rp.add_run('Reference: '), bold=True)
            for i, ref in enumerate(refs):
                if i > 0:
                    rp.add_run().add_break()
                style_run(rp.add_run(ref))

        # Light separator before the next finding
        sep_p = doc.add_paragraph()
        sep_p.paragraph_format.space_before = Pt(6)
        sep_p.paragraph_format.space_after = Pt(12)
        pPr = sep_p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'BFBFBF')
        pBdr.append(bottom)
        pPr.append(pBdr)

        doc.add_page_break()

    # ============================================================
    # 5. WEB OWASP TOP 10
    # ============================================================
    doc.add_heading('5. Web OWASP Top 10', level=1)
    owasp_data = [
        ('OWASP Top 10-2025', 'Performed'),
        ('A01:2025 - Broken Access Control', '\u2713'),
        ('A02:2025 - Security Misconfiguration', '\u2713'),
        ('A03:2025 - Software Supply Chain Failures', 'N/A'),
        ('A04:2025 - Cryptographic Failures', '\u2713'),
        ('A05:2025 - Injection', '\u2713'),
        ('A06:2025 - Insecure Design', '\u2713'),
        ('A07:2025 - Authentication Failures', 'N/A'),
        ('A08:2025 - Software or Data Integrity Failures', 'N/A'),
        ('A09:2025 - Security Logging and Alerting Failures', 'N/A'),
        ('A10:2025 - Mishandling of Exceptional Conditions', '\u2713'),
    ]
    table = doc.add_table(rows=len(owasp_data), cols=2)
    for i, (label, status) in enumerate(owasp_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = status
        table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table, header_rows=1)

    # ============================================================
    # 6. APPENDIX
    # ============================================================
    doc.add_page_break()
    doc.add_heading('6. Appendix', level=1)
    p = doc.add_paragraph()
    style_run(p.add_run('Appendix section - Non-vulnerable test cases'), italic=True)

    return doc

# ============================================================
# 7. ROUTES
# ============================================================

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/get_template/<vuln_name>')
def get_template(vuln_name):
    # Case-insensitive lookup — a tester typing "misconfigured csp header"
    # should still find the template saved as "Misconfigured CSP Header".
    template = VulnerabilityTemplate.query.filter(
        db.func.lower(VulnerabilityTemplate.name) == vuln_name.strip().lower()
    ).first()
    if template:
        return jsonify(template.to_dict())

    for t in TEMPLATES:
        if t['name'].lower() == vuln_name.strip().lower():
            return jsonify(t)

    return jsonify({'error': 'Template not found'}), 404

@app.route('/search_templates')
def search_templates():
    """Live-suggestion search used by the Vulnerability Name autocomplete
    dropdown. Searches both the growing VulnerabilityTemplate knowledge base
    and the built-in TEMPLATES starter list, case-insensitively, and ranks
    names that *start with* the query above names that merely *contain* it
    (so typing "csp" puts "CSP ..." style names above anything that just
    happens to mention csp mid-name)."""
    query = request.args.get('q', '').strip().lower()
    if not query:
        return jsonify([])

    db_templates = VulnerabilityTemplate.query.filter(
        VulnerabilityTemplate.name.ilike(f'%{query}%')
    ).all()

    results = [t.to_dict() for t in db_templates]
    seen_lower = {r['name'].strip().lower() for r in results}

    for t in TEMPLATES:
        if query in t['name'].lower() and t['name'].strip().lower() not in seen_lower:
            results.append(t)
            seen_lower.add(t['name'].strip().lower())

    results.sort(key=lambda r: (not r['name'].lower().startswith(query), r['name'].lower()))

    return jsonify(results[:8])

@app.route('/add_vulnerability', methods=['POST'])
def add_vulnerability():
    try:
        company_name = request.form.get('company_name')
        vuln_name = request.form.get('vuln_name')
        severity = request.form.get('severity')
        cwe_cve = request.form.get('cwe_cve')
        description = request.form.get('description')
        business_impact = request.form.get('business_impact')
        remediation = request.form.get('remediation')
        references = request.form.get('references')
        vulnerable_urls = request.form.get('vulnerable_urls')
        poc_steps_json = request.form.get('poc_steps')
        status = request.form.get('status', 'Open')

        images = []
        if 'poc_images' in request.files:
            files = request.files.getlist('poc_images')
            for file in files:
                if file and file.filename:
                    img_path = save_poc_image(company_name, vuln_name, file)
                    images.append(img_path)

        poc_steps = json.loads(poc_steps_json) if poc_steps_json else []
        for i, step in enumerate(poc_steps):
            if i < len(images):
                step['image'] = images[i]

        report = VulnerabilityReport(
            company_name=company_name,
            name=vuln_name,
            severity=severity,
            cwe_cve=cwe_cve,
            description=description,
            business_impact=business_impact,
            remediation=remediation,
            references=references,
            vulnerable_urls=vulnerable_urls,
            poc_steps=json.dumps(poc_steps),
            image_paths=json.dumps(images),
            status=status
        )

        db.session.add(report)

        # === NEW: Auto Vulnerability Knowledge Base ===
        # Every save teaches the template DB: first time this vuln name is
        # seen, a brand-new reusable template is created; if it already
        # exists, it's refreshed with whatever fresh detail was just typed.
        # This is what makes "type the same vuln name again next week, for a
        # different client, and have everything auto-fill" work.
        template, template_created = upsert_vulnerability_template(
            vuln_name, severity, cwe_cve, description, business_impact,
            remediation, references
        )

        db.session.commit()

        return jsonify({
            'status': 'success',
            'id': report.id,
            'template_created': template_created,
            'template_name': template.name if template else vuln_name,
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/generate_report', methods=['POST'])
def generate_report():
    # === MODIFIED: header logo & watermark are now ALWAYS the bundled
    # ProTechmanize branding (static/branding/header_logo_default.png and
    # watermark_default.png) — applied automatically to every report. This
    # is no longer something the tester uploads per-report; it's fixed
    # company branding. (An optional 'header_logo' / 'watermark_logo' file
    # can still be posted to override this for a one-off report, but the
    # UI no longer exposes that — the bundled default is always used.)
    #
    # This route also accepts:
    #   - client_name        (text field, shown in header + cover page)
    #   - report_date        (optional text field, defaults to today)
    #   - cert_in_text       (optional text field)
    #   - company_website    (optional text field)
    try:
        company_name = request.form.get('company_name')
        tester_name = request.form.get('tester_name', 'Tester')
        reviewer_name = request.form.get('reviewer_name', 'Reviewer')
        approver_name = request.form.get('approver_name', 'Approver')

        # === NEW fields ===
        client_name = request.form.get('client_name') or company_name
        report_date = request.form.get('report_date') or datetime.now().strftime('%d-%m-%Y')
        cert_in_text = request.form.get('cert_in_text') or DEFAULT_CERT_TEXT
        company_website = request.form.get('company_website') or DEFAULT_WEBSITE
        classification_text = request.form.get('classification_text') or DEFAULT_CLASSIFICATION
        address_text = request.form.get('address_text') or DEFAULT_ADDRESS

        # Bundled ProTechmanize branding is used automatically by default —
        # there's no "upload a logo/watermark" step in the UI anymore.
        # header_logo / watermark_logo are still accepted here (for a
        # one-off white-label override via direct API call), but nothing in
        # the form posts them under normal use.
        header_logo_file = request.files.get('header_logo')
        watermark_logo_file = request.files.get('watermark_logo')
        header_logo_path = save_or_reuse_asset(
            company_name, header_logo_file, 'header_logo',
            default_path=DEFAULT_HEADER_LOGO_PATH
        )
        watermark_logo_path = save_or_reuse_asset(
            company_name, watermark_logo_file, 'watermark_logo',
            default_path=DEFAULT_WATERMARK_PATH
        )

        vulns = VulnerabilityReport.query.filter_by(company_name=company_name).all()

        if not vulns:
            return jsonify({'status': 'error', 'message': 'No vulnerabilities found'}), 400

        scope_urls = [
            {'name': 'Application', 'url': request.form.get('scope_url', ''),
             'env': request.form.get('scope_env', 'UAT/PROD'),
             'methodology': request.form.get('scope_methodology', 'BlackBox')}
        ]

        vulnerabilities = [v.to_dict() for v in vulns]
        doc = generate_vapt_report(
            company_name, scope_urls, tester_name, reviewer_name, approver_name,
            vulnerabilities, client_name=client_name, header_logo_path=header_logo_path,
            watermark_logo_path=watermark_logo_path, report_date=report_date,
            cert_in_text=cert_in_text, company_website=company_website,
            classification_text=classification_text, address_text=address_text
        )

        filename = f"VAPT_Report_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = DESKTOP / filename
        doc.save(str(filepath))

        return send_file(str(filepath), as_attachment=True, download_name=filename)
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/get_vulnerabilities/<company_name>')
def get_vulnerabilities(company_name):
    vulns = VulnerabilityReport.query.filter_by(company_name=company_name).all()
    return jsonify([v.to_dict() for v in vulns])

# ============================================================
# 8. INIT DATABASE
# ============================================================

with app.app_context():
    db.create_all()

    for template in TEMPLATES:
        existing = VulnerabilityTemplate.query.filter_by(name=template['name']).first()
        if not existing:
            new_template = VulnerabilityTemplate(
                name=template['name'],
                severity=template['severity'],
                cwe_cve=template['cwe_cve'],
                description=template['description'],
                business_impact=template['business_impact'],
                remediation=template['remediation'],
                references=template['references']
            )
            db.session.add(new_template)
    db.session.commit()

# ============================================================
# 9. RUN
# ============================================================

if __name__ == '__main__':
    print("\n" + "="*50)
    print("  VAPT One-Click Report Generator")
    print("  Server: http://127.0.0.1:5000")
    print("="*50 + "\n")
    app.run(debug=True, host='127.0.0.1', port=5000)
